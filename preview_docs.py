# -*- coding: utf-8 -*-
"""
Generates CV, Cover Letter, and Motivation Letter as .docx files.
ATS-safe: Arial font, single-column body, no special Unicode chars,
no em/en dashes, no decorative bullets, no tables (except borderless header).
European executive standard: photo top-right, Austrian address, EU work auth.
"""

import os
import sys
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm

from templates.master_cv import (
    CORE_COMPETENCIES, EXPERIENCE, EDUCATION, LANGUAGES, SOFTWARE, PROFILE, SUMMARY
)
from templates.master_cover_letter import STRUCTURE

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'output')
PHOTO_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'data', 'photo.jpg')
os.makedirs(OUTPUT_DIR, exist_ok=True)

# Colors
NAVY   = RGBColor(0x1E, 0x3A, 0x5F)
BLACK  = RGBColor(0x1A, 0x1A, 0x1A)
GRAY   = RGBColor(0x55, 0x65, 0x81)
LGRAY  = RGBColor(0xA0, 0xAE, 0xC0)

SAMPLE_JOB = {
    'title': 'Director Consumer Insights',
    'company': 'Nestle Austria',
    'location': 'Vienna, Austria',
    'url': 'https://linkedin.com/jobs/view/sample',
}

# Shared "What I bring" bullets for motivation letters
MOTIVATION_BULLETS = [
    'Analytical depth. I have worked inside Nielsen and Kantar, I know how the data is built, '
    'where its limits are, and how to squeeze commercial signal out of it that most teams miss.',
    'Commercial leadership. I have sat in AOP rooms, held P&L accountability at Walmart, and '
    'built category frameworks that directly shaped $500M+ planning cycles.',
    'Organizational impact. At PepsiCo I managed teams of 12 across 8 countries and I am '
    'currently designing the operating model that determines how commercial capabilities '
    'are delivered across anchor markets globally. I know how to move organizations.',
]


# ─────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────

def set_margins(doc, top=Cm(2.0), bottom=Cm(2.0), left=Cm(2.2), right=Cm(2.2)):
    for s in doc.sections:
        s.top_margin    = top
        s.bottom_margin = bottom
        s.left_margin   = left
        s.right_margin  = right


def af(run, size=11, bold=False, italic=False, color=None):
    """Apply Arial font to a run."""
    run.font.name  = 'Arial'
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    run.font.color.rgb = color if color else BLACK


def remove_table_borders(table):
    """Remove all visible borders from a table."""
    tbl = table._tbl
    # Get or create tblPr
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    # Remove existing tblBorders if any
    for existing in tblPr.findall(qn('w:tblBorders')):
        tblPr.remove(existing)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')
        border.set(qn('w:sz'), '0')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def add_section_header(doc, text):
    """Section header: bold navy Arial 11pt + thin gray bottom border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text.upper())
    af(r, size=10.5, bold=True, color=NAVY)

    # Bottom border on the paragraph (acts as a divider line)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'A0AEC0')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_bullet(doc, text, indent_cm=0.4):
    """Plain bullet point using standard hyphen, Arial 10.5pt."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent  = Cm(indent_cm)
    r = p.add_run('- ' + text)
    af(r, size=10.5)
    return p


# ─────────────────────────────────────────
# 1. CV
# ─────────────────────────────────────────

def build_cv():
    doc = Document()
    set_margins(doc)

    # ── HEADER: borderless 2-cell table ──
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(tbl)

    # Column widths: ~75% name/contact, ~25% photo
    tbl.columns[0].width = Cm(12.5)
    tbl.columns[1].width = Cm(4.0)

    left_cell  = tbl.cell(0, 0)
    right_cell = tbl.cell(0, 1)
    right_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # Left: Name + title + contact
    lp = left_cell.paragraphs[0]
    lp.paragraph_format.space_after = Pt(2)
    r = lp.add_run('Emmanuel Rodriguez')
    r.font.name = 'Arial'
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = NAVY

    def lp_add(text, size=10.5, bold=False, color=None, space_after=2):
        p = left_cell.add_paragraph()
        p.paragraph_format.space_after  = Pt(space_after)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(text)
        af(run, size=size, bold=bold, color=color or GRAY)
        return p

    lp_add('Commercial and Insights Leader  |  Business Transformation  |  FMCG / CPG',
           size=10.5, bold=False, color=GRAY)
    lp_add('')
    lp_add('Lilienthalgasse 3, 1030 Wien, Austria',  size=10)
    lp_add('+43 66021 64853  |  emmanuel.rdrlp@gmail.com', size=10)

    # Right: Photo
    rp = right_cell.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if os.path.exists(PHOTO_PATH):
        run = rp.add_run()
        run.add_picture(PHOTO_PATH, width=Cm(3.2))
    else:
        run = rp.add_run('[Photo]')
        af(run, size=9, color=LGRAY)
        rp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # spacing

    # ── PROFESSIONAL SUMMARY ──
    add_section_header(doc, 'Professional Summary')
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)
    r = sp.add_run(SUMMARY)
    af(r, size=10.5)

    # ── SKILLS ──
    add_section_header(doc, 'Skills')
    skill_rows = [
        CORE_COMPETENCIES[0:4],
        CORE_COMPETENCIES[4:8],
        CORE_COMPETENCIES[8:12],
        CORE_COMPETENCIES[12:16],
        CORE_COMPETENCIES[16:20],
    ]
    for row in skill_rows:
        if row:
            p = doc.add_paragraph()
            p.paragraph_format.space_after  = Pt(2)
            p.paragraph_format.space_before = Pt(0)
            r = p.add_run('  /  '.join(row))
            af(r, size=10.5)

    # ── PROFESSIONAL EXPERIENCE ──
    add_section_header(doc, 'Professional Experience')

    for exp in EXPERIENCE:
        # Role title / Company / Location / Dates
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(1)
        r1 = p.add_run(exp['title'])
        af(r1, size=11, bold=True, color=BLACK)
        r2 = p.add_run(f"  /  {exp['company']}")
        af(r2, size=10.5, color=GRAY)

        dp = doc.add_paragraph()
        dp.paragraph_format.space_after  = Pt(3)
        dp.paragraph_format.space_before = Pt(0)
        r = dp.add_run(f"{exp['start']} - {exp['end']}")
        af(r, size=9.5, italic=True, color=LGRAY)

        for bullet in exp['bullets'][:3]:
            # Clean any stray dashes/special chars from bullet text
            clean = bullet.replace('\u2014', '-').replace('\u2013', '-').replace('\u2022', '')
            add_bullet(doc, clean)

    # ── EDUCATION ──
    add_section_header(doc, 'Education')
    for edu in EDUCATION:
        ep = doc.add_paragraph()
        ep.paragraph_format.space_after  = Pt(3)
        ep.paragraph_format.space_before = Pt(6)
        r1 = ep.add_run(edu['degree'])
        af(r1, size=10.5, bold=True, color=BLACK)
        r2 = ep.add_run(f"  /  {edu['institution']}, {edu['location']}")
        af(r2, size=10.5, color=GRAY)
        date_range = edu['end'] if edu['start'] == edu['end'] else f"{edu['start']} - {edu['end']}"
        r3 = ep.add_run(f"  ({date_range})")
        af(r3, size=10, italic=True, color=LGRAY)

    # ── LANGUAGES ──
    add_section_header(doc, 'Languages')
    lp2 = doc.add_paragraph()
    lp2.paragraph_format.space_after = Pt(3)
    r = lp2.add_run('  /  '.join(f"{l['language']} ({l['level']})" for l in LANGUAGES))
    af(r, size=10.5)

    # ── TOOLS ──
    add_section_header(doc, 'Tools and Software')
    tp = doc.add_paragraph()
    r = tp.add_run('  /  '.join(SOFTWARE))
    af(r, size=10.5)

    path = os.path.join(OUTPUT_DIR, 'CV_Emmanuel_Rodriguez_v2.docx')
    doc.save(path)
    print(f'[OK] CV saved: {path}')
    return path


# ─────────────────────────────────────────
# 2. COVER LETTER
# ─────────────────────────────────────────

def build_cover_letter():
    doc = Document()
    set_margins(doc, top=Cm(2.5), bottom=Cm(2.5), left=Cm(2.8), right=Cm(2.8))

    def ap(text, size=11, bold=False, color=None, space_after=8, space_before=0, align=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(space_after)
        p.paragraph_format.space_before = Pt(space_before)
        if align:
            p.alignment = align
        r = p.add_run(text)
        af(r, size=size, bold=bold, color=color or BLACK)
        return p

    # Header
    ap('Emmanuel Rodriguez', size=16, bold=True, color=NAVY, space_after=2)
    ap('Lilienthalgasse 3, 1030 Wien, Austria', size=10, color=GRAY, space_after=1)
    ap('emmanuel.rdrlp@gmail.com  /  +43 66021 64853  /  linkedin.com/in/emmanuel-rdrlp/',
       size=10, color=GRAY, space_after=16)

    ap(datetime.now().strftime('%B %d, %Y'), size=11, space_after=14)

    ap(f"Hiring Manager\n{SAMPLE_JOB['company']}\n{SAMPLE_JOB['location']}",
       size=11, space_after=14)

    ap(f"Re: {SAMPLE_JOB['title']}",
       size=11, bold=True, color=NAVY, space_after=14)

    ap('Dear Hiring Team,', size=11, space_after=10)

    # Opening - no "I am writing to express..."
    ap(
        f"The {SAMPLE_JOB['title']} role at {SAMPLE_JOB['company']} fits the arc of my career "
        f"precisely. Ten years of building consumer insights and analytics capabilities across "
        f"FMCG markets - PepsiCo, Nielsen, Kantar - have given me a sharp view of where data "
        f"creates commercial advantage and where it creates noise. I know the difference.",
        size=11, space_after=10
    )

    # Body - specific, no buzzwords
    ap(
        f"At PepsiCo LATAM, I led the Consumer and Market Insights function for a business of "
        f"roughly $2 billion in annual revenue. The work went well beyond reporting: I built the "
        f"analytics architecture from the ground up, cutting time-to-insight by 40%, and sat in "
        f"the room where annual operating plan decisions were made. Before that, as Analytics Manager "
        f"I ran a team of 12 across 8 countries, standardizing how the region read the market. "
        f"In my current role in Business Transformation, I design the commercial operating models "
        f"that determine how those insights actually reach the frontline.",
        size=11, space_after=10
    )

    # Company fit
    ap(
        f"{SAMPLE_JOB['company']}'s position in the Austrian and European consumer goods market, "
        f"and its track record of using category data to drive shelf performance, is exactly "
        f"the environment where this type of work produces results. I relocated to Austria "
        f"with a clear purpose and I am currently completing my MBA at TU Wien.",
        size=11, space_after=10
    )

    # Work auth - short, factual
    ap(STRUCTURE['austria_paragraph'], size=11, space_after=10)

    # Closing - direct
    ap(
        f"I would welcome a conversation about what the team is building. Thank you for your time.",
        size=11, space_after=20
    )

    ap('Kind regards,', size=11, space_after=4)
    ap('Emmanuel Rodriguez', size=11, bold=True, space_after=2)

    path = os.path.join(OUTPUT_DIR, 'CoverLetter_Emmanuel_Rodriguez_v2.docx')
    doc.save(path)
    print(f'[OK] Cover Letter saved: {path}')
    return path


def build_master_cover_letter():
    """Standalone, general-purpose cover letter (counterpart to the Master CV).
    Fill in [Position Title] and [Company Name] before sending."""
    doc = Document()
    set_margins(doc, top=Cm(2.5), bottom=Cm(2.5), left=Cm(2.8), right=Cm(2.8))

    def ap(text, size=11, bold=False, color=None, space_after=8, space_before=0, align=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(space_after)
        p.paragraph_format.space_before = Pt(space_before)
        if align:
            p.alignment = align
        r = p.add_run(text)
        af(r, size=size, bold=bold, color=color or BLACK)
        return p

    # Header
    ap('Emmanuel Rodriguez', size=16, bold=True, color=NAVY, space_after=2)
    ap(PROFILE['address'], size=10, color=GRAY, space_after=1)
    ap(f"{PROFILE['email']}  /  {PROFILE['phone']}  /  linkedin.com/in/emmanuel-rdrlp/",
       size=10, color=GRAY, space_after=16)

    ap(datetime.now().strftime('%B %d, %Y'), size=11, space_after=14)

    ap('Re: [Position Title] - [Company Name]', size=11, bold=True, color=NAVY, space_after=14)

    ap('Dear Hiring Manager,', size=11, space_after=10)

    ap(
        "The [Position Title] role at [Company Name] sits at the intersection of where my "
        "career has been heading for the past decade: turning consumer and market data into "
        "commercial strategy, and now into the operating models that make that strategy "
        "executable at scale.",
        size=11, space_after=10
    )

    ap(
        "Over 10+ years across FMCG - Nielsen, Kantar, Walmart and PepsiCo - I have moved from "
        "producing market insight to owning the strategy built on it, and now to leading the "
        "organizational change behind it. As Analytics & Insights Manager and later SR Manager, "
        "Market & Consumer Insights at PepsiCo, I led regional teams of 12+ professionals, "
        "supported $500M+ annual planning cycles, and cut time-to-insight by 40%. In my current "
        "role as Business Transformation Lead, Commercial & Sales at PepsiCo Global, I am "
        "redesigning operating models across anchor markets, reducing operational complexity by "
        "around 30% while coordinating 150+ stakeholders across Commercial, Sales, Finance and "
        "Capabilities.",
        size=11, space_after=10
    )

    ap(
        "Based in Vienna and completing an MBA in Advanced Technologies & Global Leadership at "
        "TU Wien, my current mandate already spans anchor markets across the Americas, Asia and "
        "Europe - and [Company Name]'s position in [industry / market] is exactly the kind of "
        "environment where that global commercial and transformation experience creates value "
        "quickly.",
        size=11, space_after=10
    )

    ap(STRUCTURE['austria_paragraph'], size=11, space_after=10)

    ap(
        "I would welcome the opportunity to discuss how this experience can contribute to "
        "[Company Name]'s goals. Thank you for your time and consideration.",
        size=11, space_after=20
    )

    ap('Kind regards,', size=11, space_after=4)
    ap('Emmanuel Rodriguez', size=11, bold=True, space_after=2)

    path = os.path.join(OUTPUT_DIR, 'Master_Cover_Letter_Emmanuel_Rodriguez.docx')
    doc.save(path)
    print(f'[OK] Master Cover Letter saved: {path}')
    return path


# ─────────────────────────────────────────
# 3. MOTIVATION LETTER
# ─────────────────────────────────────────

def build_motivation_letter():
    doc = Document()
    set_margins(doc, top=Cm(2.5), bottom=Cm(2.5), left=Cm(2.8), right=Cm(2.8))

    def ap(text, size=11, bold=False, color=None, space_after=8, space_before=0):
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(space_after)
        p.paragraph_format.space_before = Pt(space_before)
        r = p.add_run(text)
        af(r, size=size, bold=bold, color=color or BLACK)
        return p

    # Header
    ap('Emmanuel Rodriguez', size=16, bold=True, color=NAVY, space_after=2)
    ap('Lilienthalgasse 3, 1030 Wien, Austria', size=10, color=GRAY, space_after=1)
    ap('emmanuel.rdrlp@gmail.com  /  +43 66021 64853', size=10, color=GRAY, space_after=16)
    ap(datetime.now().strftime('%B %d, %Y'), size=11, space_after=14)
    ap(f"Motivation Letter\n{SAMPLE_JOB['title']}  /  {SAMPLE_JOB['company']}",
       size=12, bold=True, color=NAVY, space_after=16)
    ap('Dear Hiring Team,', size=11, space_after=10)

    # Section 1: Why this role
    ap('Why this role', size=11, bold=True, color=NAVY, space_after=4)
    ap(
        f"The {SAMPLE_JOB['title']} position at {SAMPLE_JOB['company']} is the role I have been "
        f"building toward. My entire career sits at the intersection of consumer data and "
        f"commercial decisions: I started as a market analyst at Nielsen reading FMCG shelf data, "
        f"moved to Kantar running custom research for category teams, then spent six years at "
        f"PepsiCo building and scaling the insights function across Latin America. "
        f"The companies that win in FMCG are the ones that treat insights as a competitive weapon, "
        f"not a reporting obligation. That is the standard I hold myself to and what I want to "
        f"help build here.",
        size=11, space_after=10
    )

    # Section 2: What I bring
    ap('What I bring', size=11, bold=True, color=NAVY, space_after=4)
    ap(
        'Three capabilities set my profile apart for this role:',
        size=11, space_after=4
    )
    for b in MOTIVATION_BULLETS:
        add_bullet(doc, b, indent_cm=0.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Section 3: Why this company and Austria
    ap(f"Why {SAMPLE_JOB['company']} and why Austria", size=11, bold=True, color=NAVY, space_after=4)
    ap(
        f"I moved to Austria with a deliberate plan. I am studying for my MBA at TU Wien, "
        f"focused on how European business models differ from the high-growth emerging markets "
        f"where I built my career. {SAMPLE_JOB['company']}'s presence in the Austrian market and "
        f"its reputation for category discipline make it the right company to apply what I know "
        f"and learn what I do not yet know about European consumer behavior.",
        size=11, space_after=10
    )
    ap(STRUCTURE['austria_paragraph'], size=11, space_after=10)

    # Section 4: First 90 days - concrete, not generic
    ap('First 90 days', size=11, bold=True, color=NAVY, space_after=4)
    days = [
        'Weeks 1 to 4: understand the current data infrastructure, the reporting cadence, '
        'and which decisions the commercial teams are actually making without good information.',
        'Weeks 5 to 8: identify the two or three insight gaps creating the most commercial risk '
        'and propose fast, practical solutions that do not require a 6-month project.',
        'Weeks 9 to 12: present a 12-month insights capability roadmap to leadership '
        'with clear owners, milestones, and business cases.',
    ]
    for d in days:
        add_bullet(doc, d, indent_cm=0.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    ap(
        'Thank you for taking the time to review my application. '
        'I look forward to the opportunity to speak.',
        size=11, space_after=20
    )
    ap('Kind regards,', size=11, space_after=4)
    ap('Emmanuel Rodriguez', size=11, bold=True)

    path = os.path.join(OUTPUT_DIR, 'MotivationLetter_Emmanuel_Rodriguez_v2.docx')
    doc.save(path)
    print(f'[OK] Motivation Letter saved: {path}')
    return path


# ─────────────────────────────────────────
# 4. GENERIC COVER LETTER (no placeholders)
# ─────────────────────────────────────────

def build_generic_cover_letter():
    """Standalone, fully self-contained cover letter for general/networking use.
    No [Position Title] / [Company Name] placeholders."""
    doc = Document()
    set_margins(doc, top=Cm(2.5), bottom=Cm(2.5), left=Cm(2.8), right=Cm(2.8))

    def ap(text, size=11, bold=False, color=None, space_after=8, space_before=0, align=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(space_after)
        p.paragraph_format.space_before = Pt(space_before)
        if align:
            p.alignment = align
        r = p.add_run(text)
        af(r, size=size, bold=bold, color=color or BLACK)
        return p

    # Header
    ap('Emmanuel Rodriguez', size=16, bold=True, color=NAVY, space_after=2)
    ap(PROFILE['address'], size=10, color=GRAY, space_after=1)
    ap(f"{PROFILE['email']}  /  {PROFILE['phone']}  /  linkedin.com/in/emmanuel-rdrlp/",
       size=10, color=GRAY, space_after=16)

    ap(datetime.now().strftime('%B %d, %Y'), size=11, space_after=14)

    ap('Dear Hiring Manager,', size=11, space_after=10)

    ap(
        "I am reaching out to explore opportunities in Corporate Strategy, Business "
        "Transformation, and Commercial & Insights Leadership within FMCG and consumer-facing "
        "organizations. My career has moved from producing market insight, to owning the "
        "strategy built on it, to leading the organizational change behind it.",
        size=11, space_after=10
    )

    ap(
        "Across 10+ years at Nielsen, Kantar, Walmart and PepsiCo, I have led regional teams of "
        "12+ professionals, supported $500M+ annual planning cycles, and cut time-to-insight by "
        "40%. In my current role as Business Transformation Lead, Commercial & Sales at PepsiCo "
        "Global, I am redesigning operating models across anchor markets, reducing operational "
        "complexity by around 30% while coordinating 150+ stakeholders across Commercial, Sales, "
        "Finance and Capabilities.",
        size=11, space_after=10
    )

    ap(
        "Based in Vienna and completing an MBA in Advanced Technologies & Global Leadership at "
        "TU Wien, my current mandate already spans anchor markets across the Americas, Asia and "
        "Europe, and I am looking for an organization where that global commercial and "
        "transformation experience, paired with deep analytical foundations, can create value "
        "quickly.",
        size=11, space_after=10
    )

    ap(STRUCTURE['austria_paragraph'], size=11, space_after=10)

    ap(
        "I would welcome the opportunity to discuss how my background could support your team's "
        "goals. Thank you for your time and consideration.",
        size=11, space_after=20
    )

    ap('Kind regards,', size=11, space_after=4)
    ap('Emmanuel Rodriguez', size=11, bold=True, space_after=2)

    path = os.path.join(OUTPUT_DIR, 'Generic_Cover_Letter_Emmanuel_Rodriguez.docx')
    doc.save(path)
    print(f'[OK] Generic Cover Letter saved: {path}')
    return path


# ─────────────────────────────────────────
# 5. MASTER MOTIVATION LETTER (with placeholders)
# ─────────────────────────────────────────

def build_master_motivation_letter():
    """Standalone motivation letter template (counterpart to the Master Cover Letter).
    Fill in [Position Title] and [Company Name] before sending."""
    doc = Document()
    set_margins(doc, top=Cm(2.5), bottom=Cm(2.5), left=Cm(2.8), right=Cm(2.8))

    def ap(text, size=11, bold=False, color=None, space_after=8, space_before=0):
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(space_after)
        p.paragraph_format.space_before = Pt(space_before)
        r = p.add_run(text)
        af(r, size=size, bold=bold, color=color or BLACK)
        return p

    # Header
    ap('Emmanuel Rodriguez', size=16, bold=True, color=NAVY, space_after=2)
    ap(PROFILE['address'], size=10, color=GRAY, space_after=1)
    ap(f"{PROFILE['email']}  /  {PROFILE['phone']}", size=10, color=GRAY, space_after=16)
    ap(datetime.now().strftime('%B %d, %Y'), size=11, space_after=14)
    ap('Motivation Letter\n[Position Title]  /  [Company Name]',
       size=12, bold=True, color=NAVY, space_after=16)
    ap('Dear Hiring Team,', size=11, space_after=10)

    # Section 1: Why this role
    ap('Why this role', size=11, bold=True, color=NAVY, space_after=4)
    ap(
        "The [Position Title] position at [Company Name] is the kind of role I have been "
        "building toward. My career sits at the intersection of consumer data and commercial "
        "decisions: I started as a market analyst at Nielsen reading FMCG shelf data, moved to "
        "Kantar running custom research for category teams, then spent six years at PepsiCo "
        "scaling the insights function across Latin America before moving into Business "
        "Transformation, where I now design the operating models that turn strategy into "
        "execution. [Company Name]'s position in [industry / market] is exactly the kind of "
        "environment where this combination of insight, strategy and execution creates value.",
        size=11, space_after=10
    )

    # Section 2: What I bring
    ap('What I bring', size=11, bold=True, color=NAVY, space_after=4)
    ap('Three capabilities set my profile apart for this role:', size=11, space_after=4)
    for b in MOTIVATION_BULLETS:
        add_bullet(doc, b, indent_cm=0.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Section 3: Why [Company Name] and why Austria
    ap('Why [Company Name] and why Austria', size=11, bold=True, color=NAVY, space_after=4)
    ap(
        "I moved to Austria with a deliberate plan: completing my MBA in Advanced Technologies "
        "& Global Leadership at TU Wien while based at the center of my current global mandate, "
        "which already spans anchor markets across the Americas, Asia and Europe. [Company "
        "Name]'s [reputation / position in industry] makes it a company where that global "
        "commercial and transformation experience can create value from day one.",
        size=11, space_after=10
    )
    ap(STRUCTURE['austria_paragraph'], size=11, space_after=10)

    # Section 4: First 90 days
    ap('First 90 days', size=11, bold=True, color=NAVY, space_after=4)
    days = [
        'Weeks 1 to 4: map the current operating model, data and decision-making landscape, '
        'and identify which decisions are being made without the right information or structure.',
        'Weeks 5 to 8: identify the two or three highest-impact opportunities and build the '
        'business case and quick wins around them.',
        'Weeks 9 to 12: present a roadmap to leadership with clear owners, milestones and '
        'measurable outcomes.',
    ]
    for d in days:
        add_bullet(doc, d, indent_cm=0.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    ap(
        'Thank you for taking the time to review my application. '
        'I look forward to the opportunity to speak.',
        size=11, space_after=20
    )
    ap('Kind regards,', size=11, space_after=4)
    ap('Emmanuel Rodriguez', size=11, bold=True)

    path = os.path.join(OUTPUT_DIR, 'Master_Motivation_Letter_Emmanuel_Rodriguez.docx')
    doc.save(path)
    print(f'[OK] Master Motivation Letter saved: {path}')
    return path


# ─────────────────────────────────────────
# 6. GENERIC MOTIVATION LETTER (no placeholders)
# ─────────────────────────────────────────

def build_generic_motivation_letter():
    """Standalone, fully self-contained motivation letter for general/networking use.
    No [Position Title] / [Company Name] placeholders."""
    doc = Document()
    set_margins(doc, top=Cm(2.5), bottom=Cm(2.5), left=Cm(2.8), right=Cm(2.8))

    def ap(text, size=11, bold=False, color=None, space_after=8, space_before=0):
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(space_after)
        p.paragraph_format.space_before = Pt(space_before)
        r = p.add_run(text)
        af(r, size=size, bold=bold, color=color or BLACK)
        return p

    # Header
    ap('Emmanuel Rodriguez', size=16, bold=True, color=NAVY, space_after=2)
    ap(PROFILE['address'], size=10, color=GRAY, space_after=1)
    ap(f"{PROFILE['email']}  /  {PROFILE['phone']}", size=10, color=GRAY, space_after=16)
    ap(datetime.now().strftime('%B %d, %Y'), size=11, space_after=14)
    ap('Motivation Letter', size=12, bold=True, color=NAVY, space_after=16)
    ap('Dear Hiring Team,', size=11, space_after=10)

    # Section 1: What drives me
    ap('What drives me', size=11, bold=True, color=NAVY, space_after=4)
    ap(
        "My career sits at the intersection of consumer data and commercial decisions. I started "
        "as a market analyst at Nielsen reading FMCG shelf data, moved to Kantar running custom "
        "research for category teams, then spent six years at PepsiCo scaling the insights "
        "function across Latin America before moving into Business Transformation, where I now "
        "design the operating models that turn strategy into execution. The organizations that "
        "win in FMCG and consumer-facing industries are the ones that treat insight and operating "
        "discipline as competitive weapons, not reporting obligations. That is the standard I "
        "hold myself to, and the kind of environment I am looking to join next.",
        size=11, space_after=10
    )

    # Section 2: What I bring
    ap('What I bring', size=11, bold=True, color=NAVY, space_after=4)
    ap('Three capabilities set my profile apart:', size=11, space_after=4)
    for b in MOTIVATION_BULLETS:
        add_bullet(doc, b, indent_cm=0.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Section 3: Why Austria and Europe
    ap('Why Austria and Europe', size=11, bold=True, color=NAVY, space_after=4)
    ap(
        "I moved to Austria with a deliberate plan: completing my MBA in Advanced Technologies "
        "& Global Leadership at TU Wien while based at the center of my current global mandate, "
        "which already spans anchor markets across the Americas, Asia and Europe. I am looking "
        "to bring that global commercial and transformation experience into a European "
        "organization where it can create value from day one.",
        size=11, space_after=10
    )
    ap(STRUCTURE['austria_paragraph'], size=11, space_after=10)

    # Section 4: How I work
    ap('How I work', size=11, bold=True, color=NAVY, space_after=4)
    ap(
        "Regardless of the specific mandate, my approach follows the same pattern: understand "
        "the current operating model and the decisions people are actually making, or struggling "
        "to make, without the right information or structure; identify the two or three "
        "opportunities that would create the most value fastest; and build a roadmap with clear "
        "owners, milestones and business cases - not a deck that sits in a drawer.",
        size=11, space_after=10
    )

    ap(
        'Thank you for taking the time to review my application. '
        'I look forward to the opportunity to speak.',
        size=11, space_after=20
    )
    ap('Kind regards,', size=11, space_after=4)
    ap('Emmanuel Rodriguez', size=11, bold=True)

    path = os.path.join(OUTPUT_DIR, 'Generic_Motivation_Letter_Emmanuel_Rodriguez.docx')
    doc.save(path)
    print(f'[OK] Generic Motivation Letter saved: {path}')
    return path


if __name__ == '__main__':
    print('\nGenerating sample documents...\n')
    cv   = build_cv()
    cl   = build_cover_letter()
    ml   = build_motivation_letter()
    mcl  = build_master_cover_letter()
    gcl  = build_generic_cover_letter()
    mml  = build_master_motivation_letter()
    gml  = build_generic_motivation_letter()
    print(f'\nAll files saved to: {OUTPUT_DIR}')

    import subprocess
    subprocess.Popen(f'explorer "{OUTPUT_DIR}"', shell=True)
