"""
Claude API-powered CV and cover letter generator.
Takes a job dict + master profile data → produces tailored .docx files.
Uses claude-sonnet-4-6 to tailor content to the specific job description.
"""

import json
import os
import re
import subprocess
import sys
from datetime import datetime

import anthropic
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from dotenv import load_dotenv

load_dotenv()

sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))
from src.database import update_status
from src.mailer import send_manual_package
from src.tracker import record_sent_cv, EXCEL_PATH
from templates.master_cover_letter import HEADER, STRUCTURE
from templates.master_cv import (
    CORE_COMPETENCIES, EDUCATION, EXPERIENCE, LANGUAGES, PROFILE,
    SOFTWARE, SUMMARY,
)

OUTPUT_DIR = os.path.join(os.path.dirname(__file__), '..', 'output')
PHOTO_PATH = os.path.join(os.path.dirname(__file__), '..', 'Originales', 'Foto Emmma V2.jpg')
os.makedirs(OUTPUT_DIR, exist_ok=True)

client = anthropic.Anthropic(api_key=os.getenv('ANTHROPIC_API_KEY', ''))

CV_SYSTEM_PROMPT = """You are an elite ATS optimization consultant for senior FMCG executive roles.

You receive:
1. Emmanuel's master CV profile (JSON)
2. A job description

Your task: produce a tailored CV as JSON. Rules:
- NEVER fabricate experience or facts — only rearrange and re-emphasize what exists
- Rewrite the summary (max 3 sentences) mirroring the job's exact keywords
- ALWAYS include ALL roles from the master profile — never drop any position
- For each role, select the 3–5 most relevant bullet points (from master bullets list)
- Reorder core_competencies to put the most relevant ones first (max 12)
- Quantify every bullet — include %, dollar amounts, headcount, scale
- Executive tone: action verbs, business impact, strategic framing
- Do NOT add extra roles not in the master profile
- WRITING STYLE — human, natural, senior executive tone:
  * NO em dashes (— or –). Use a period or rewrite the sentence instead
  * NO Oxford commas (no comma before "and" in a list)
  * NO over-formal connectors like "Furthermore", "Moreover", "Additionally"
  * Vary sentence structure — not every bullet should start the same way
  * DO NOT mention visa, residency, work authorization or Austria eligibility anywhere in the CV content (it is already in the header)

Return ONLY valid JSON, no markdown, no explanation:
{
  "summary": "...",
  "core_competencies": ["...", ...],
  "experience": [
    {
      "title": "...",
      "company": "...",
      "start": "...",
      "end": "...",
      "bullets": ["...", ...]
    }
  ]
}"""

CL_SYSTEM_PROMPT = """You are an elite executive cover letter writer for senior FMCG roles.

COVER LETTER PURPOSE: A professional pitch — WHAT YOU BRING to this role.
Focus: your most relevant achievements, skills match, and concrete value-add.
This is NOT about why you want the role — that goes in the Motivation Letter.

You receive:
1. Emmanuel's profile summary
2. A job description

Your task: write a tailored cover letter. Hard rules:
- STRICT 1-PAGE LIMIT: max 3 short paragraphs, 180–210 words total (body only, excluding salutation/sign-off)
- Each paragraph max 3 sentences
- Opening (1 paragraph): name the role, state your most relevant title/expertise and ONE headline achievement with a number
- Body (1 paragraph): 2 specific accomplishments that directly mirror the job's key requirements — use metrics (%, scale, revenue)
- Closing (1 paragraph): one sentence on fit + one forward-looking sentence
- Mirror keywords from the job description naturally
- NO generic opener ("I am writing to apply…") — start with impact
- WRITING STYLE:
  * NO em dashes (— or –). Use a period or rewrite instead
  * NO Oxford commas
  * NO over-formal connectors ("Furthermore", "Moreover", "Additionally")
  * DO NOT mention visa, residency or work authorization (added separately)

Return ONLY valid JSON:
{
  "salutation": "Dear [Name/Hiring Team],",
  "opening": "...",
  "body": "...",
  "closing": "...",
  "sign_off": "Kind regards,\\nEmmanuel Rodríguez"
}"""


ML_SYSTEM_PROMPT = """You are writing a Letter of Motivation for Emmanuel Rodríguez applying to a senior FMCG role.

MOTIVATION LETTER PURPOSE: WHY you want this role and this company — personal drive, career story, cultural fit.
This is NOT a list of achievements — that belongs in the Cover Letter.

Write 3 paragraphs, 200–230 words total (body only, excluding salutation/sign-off). Structure:
1. Opening (3–4 sentences): Why THIS specific role at THIS specific company genuinely excites Emmanuel. Be concrete — reference the company's market position, product portfolio or strategic direction. Not generic enthusiasm.
2. Career narrative + field drive (3–4 sentences): The professional journey that makes this role the natural next step. Why Consumer Insights / Analytics / Commercial Strategy / Transformation is Emmanuel's field of choice — the intellectual pull, the impact seen firsthand.
3. Closing (2–3 sentences): Why this is the right moment in Emmanuel's career. Include Austria residency as a single natural sentence (not a standalone announcement). Confident forward look.

Hard rules:
- STRICT 1-PAGE LIMIT: 200–230 words body total
- NEVER fabricate — draw only from Emmanuel's real background
- Human, personal, genuine — no corporate boilerplate
- Each paragraph must feel distinct from the Cover Letter tone
- NO em dashes (— or –). Use a period or rewrite instead
- NO Oxford commas
- NO over-formal connectors ("Furthermore", "Moreover", "Additionally")

Return ONLY valid JSON:
{
  "salutation": "Dear [Name/Hiring Team],",
  "opening": "...",
  "motivation_field": "...",
  "closing": "...",
  "sign_off": "Kind regards,\\nEmmanuel Rodríguez"
}"""


def _call_claude(system: str, user_message: str, max_tokens: int = 2000) -> str:
    message = client.messages.create(
        model='claude-sonnet-4-6',
        max_tokens=max_tokens,
        system=system,
        messages=[{'role': 'user', 'content': user_message}],
    )
    return message.content[0].text


def _tailor_cv(job: dict) -> dict:
    master_data = {
        'summary': SUMMARY,
        'core_competencies': CORE_COMPETENCIES,
        'experience': EXPERIENCE,
        'profile': PROFILE,
    }
    user_msg = f"""MASTER PROFILE:
{json.dumps(master_data, ensure_ascii=False, indent=2)}

JOB DESCRIPTION:
Title: {job['title']}
Company: {job['company']}
Location: {job.get('location', '')}

{job.get('description', '')[:3000]}"""

    raw = _call_claude(CV_SYSTEM_PROMPT, user_msg)
    # Strip any accidental markdown
    raw = re.sub(r'^```json\s*', '', raw.strip())
    raw = re.sub(r'\s*```$', '', raw.strip())
    return json.loads(raw)


def _tailor_cover_letter(job: dict) -> dict:
    structure_str = json.dumps(STRUCTURE, ensure_ascii=False, indent=2)
    user_msg = f"""PROFILE SUMMARY:
{SUMMARY}

CONTACT:
{PROFILE['name']} | {PROFILE['address']} | {PROFILE['phone']} | {PROFILE['email']}
Work authorization: {PROFILE['work_authorization']}

COVER LETTER STRUCTURE TEMPLATE:
{structure_str}

JOB DESCRIPTION:
Title: {job['title']}
Company: {job['company']}
Location: {job.get('location', '')}

{job.get('description', '')[:3000]}"""

    raw = _call_claude(CL_SYSTEM_PROMPT, user_msg)
    raw = re.sub(r'^```json\s*', '', raw.strip())
    raw = re.sub(r'\s*```$', '', raw.strip())
    return json.loads(raw)


def _tailor_motivation_letter(job: dict) -> dict:
    user_msg = f"""PROFILE SUMMARY:
{SUMMARY}

CONTACT:
{PROFILE['name']} | {PROFILE['address']} | {PROFILE['phone']} | {PROFILE['email']}

JOB DESCRIPTION:
Title: {job['title']}
Company: {job['company']}
Location: {job.get('location', '')}

{job.get('description', '')[:3000]}"""

    raw = _call_claude(ML_SYSTEM_PROMPT, user_msg)
    raw = re.sub(r'^```json\s*', '', raw.strip())
    raw = re.sub(r'\s*```$', '', raw.strip())
    return json.loads(raw)


def _safe_filename(text: str) -> str:
    return re.sub(r'[^\w\-]', '_', text)[:30]


def _docx_to_pdf(docx_path: str) -> str | None:
    """Convert a .docx to .pdf using LibreOffice. Returns pdf path or None on failure."""
    import shutil
    lo = shutil.which('libreoffice') or shutil.which('soffice')
    if not lo:
        for candidate in ['/usr/bin/libreoffice', '/usr/bin/soffice',
                          '/usr/lib/libreoffice/program/soffice']:
            if os.path.exists(candidate):
                lo = candidate
                break
    if not lo:
        print(f"[generator] LibreOffice not found — sending DOCX", flush=True)
        return None
    print(f"[generator] Using LibreOffice at: {lo}", flush=True)
    try:
        out_dir = os.path.dirname(docx_path)
        print(f"[generator] Converting to PDF: {docx_path}", flush=True)
        env = {**os.environ, 'HOME': '/tmp'}
        result = subprocess.run(
            [lo, '--headless', '--norestore', '--convert-to', 'pdf', '--outdir', out_dir, docx_path],
            capture_output=True, text=True, timeout=60, env=env
        )
        print(f"[generator] LO returncode={result.returncode} stdout={result.stdout[:100]} stderr={result.stderr[:100]}", flush=True)
        if result.returncode == 0:
            pdf_path = docx_path.replace('.docx', '.pdf')
            if os.path.exists(pdf_path):
                print(f"[generator] PDF created: {pdf_path}", flush=True)
                return pdf_path
            print(f"[generator] PDF file not found after conversion", flush=True)
        else:
            print(f"[generator] PDF conversion failed rc={result.returncode}: {result.stderr[:200]}", flush=True)
    except Exception as e:
        print(f"[generator] PDF conversion error: {e}", flush=True)
    return None


# --- DOCX builders ---

def _set_font(run, size=11, bold=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def _remove_table_borders(table):
    """Remove all visible borders from a table (invisible layout table)."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _build_cv_docx(tailored: dict, job: dict) -> str:
    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(54)
        section.right_margin = Pt(54)

    # Header: name+contact on left, photo on right
    header_table = doc.add_table(rows=1, cols=2)
    _remove_table_borders(header_table)
    left_cell = header_table.cell(0, 0)
    right_cell = header_table.cell(0, 1)

    # Left: name
    name_p = left_cell.paragraphs[0]
    run = name_p.add_run(PROFILE['name'])
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 58, 95)

    # Left: contact
    contact_p = left_cell.add_paragraph()
    contact_p.paragraph_format.space_after = Pt(2)
    run = contact_p.add_run(
        f"{PROFILE['address']}  |  {PROFILE['phone']}  |  {PROFILE['email']}  |  {PROFILE['linkedin']}"
    )
    _set_font(run, size=9, color=(107, 114, 128))

    # Left: work authorization
    auth_p = left_cell.add_paragraph()
    auth_p.paragraph_format.space_after = Pt(4)
    run = auth_p.add_run(PROFILE['work_authorization'])
    _set_font(run, size=9, bold=True, color=(37, 99, 235))

    # Right: photo
    photo_p = right_cell.paragraphs[0]
    photo_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if os.path.exists(PHOTO_PATH):
        photo_run = photo_p.add_run()
        photo_run.add_picture(PHOTO_PATH, width=Cm(2.8))

    # Divider
    doc.add_paragraph('─' * 80).paragraph_format.space_after = Pt(2)

    # Summary
    _add_section_header(doc, 'PROFESSIONAL SUMMARY')
    p = doc.add_paragraph()
    run = p.add_run(tailored.get('summary', SUMMARY))
    _set_font(run, size=10.5)
    p.paragraph_format.space_after = Pt(8)

    # Core Competencies
    _add_section_header(doc, 'CORE COMPETENCIES')
    comps = tailored.get('core_competencies', CORE_COMPETENCIES)[:12]
    # 3-column layout via tab stops
    for i in range(0, len(comps), 3):
        row = comps[i:i+3]
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run('  ·  '.join(f'▪ {c}' for c in row))
        _set_font(run, size=9.5)

    doc.add_paragraph('').paragraph_format.space_after = Pt(4)

    # Experience — guarantee all master roles appear even if Claude dropped some
    tailored_exp = tailored.get('experience', [])
    tailored_titles = {e.get('title', '').lower() for e in tailored_exp}
    for master_role in EXPERIENCE:
        if master_role['title'].lower() not in tailored_titles:
            tailored_exp.append(master_role)

    # Experience
    _add_section_header(doc, 'PROFESSIONAL EXPERIENCE')
    for exp in tailored_exp:
        # Role title + company
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(exp['title'])
        _set_font(run, size=11, bold=True, color=(30, 58, 95))
        run2 = p.add_run(f"  —  {exp['company']}")
        _set_font(run2, size=10.5, color=(75, 85, 99))

        # Dates
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(2)
        run = p2.add_run(f"{exp.get('start', '')} – {exp.get('end', '')}")
        _set_font(run, size=9, color=(156, 163, 175))

        # Bullets
        for bullet in exp.get('bullets', []):
            bp = doc.add_paragraph(style='List Bullet')
            bp.paragraph_format.space_after = Pt(1)
            run = bp.add_run(bullet)
            _set_font(run, size=10)

    doc.add_paragraph('').paragraph_format.space_after = Pt(4)

    # Education
    _add_section_header(doc, 'EDUCATION')
    for edu in EDUCATION:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{edu['degree']}  —  {edu['institution']}, {edu['location']}")
        _set_font(run, size=10.5, bold=True)
        run2 = p.add_run(f"  ({edu['start']} – {edu['end']})")
        _set_font(run2, size=9.5, color=(107, 114, 128))

    # Languages + Software
    _add_section_header(doc, 'LANGUAGES & TOOLS')
    lang_str = '  |  '.join(f"{l['language']} ({l['level']})" for l in LANGUAGES)
    sw_str = '  ·  '.join(SOFTWARE)
    p = doc.add_paragraph()
    run = p.add_run(f"Languages: {lang_str}")
    _set_font(run, size=10)
    p2 = doc.add_paragraph()
    run = p2.add_run(f"Tools: {sw_str}")
    _set_font(run, size=10)

    # Save
    fname = "Emmanuel_Rodriguez_CV.docx"
    path = os.path.join(OUTPUT_DIR, fname)
    doc.save(path)
    print(f"[generator] CV saved: {path}")
    return path


def _add_section_header(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    _set_font(run, size=10, bold=True, color=(30, 58, 95))
    run2 = p.add_run('\n' + '─' * 60)
    _set_font(run2, size=8, color=(209, 213, 219))


def _build_letter_docx(tailored: dict, _job: dict, para_keys: list[str],
                        title_label: str, fname_prefix: str,
                        auth_paragraph: bool = False) -> str:
    """Shared builder for Cover Letter and Motivation Letter. Enforces 1-page layout."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Pt(42)
        section.bottom_margin = Pt(42)
        section.left_margin = Pt(60)
        section.right_margin = Pt(60)

    # Header: name
    h = doc.add_heading(HEADER['name'], 0)
    h.paragraph_format.space_after = Pt(2)
    for run in h.runs:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(30, 58, 95)

    # Contact line
    contact_p = doc.add_paragraph(
        f"{HEADER['address']}  ·  {HEADER['email']}  ·  {HEADER['phone']}"
    )
    for run in contact_p.runs:
        _set_font(run, size=9.5, color=(107, 114, 128))
    contact_p.paragraph_format.space_after = Pt(2)

    # Thin divider
    div = doc.add_paragraph('─' * 72)
    div.paragraph_format.space_after = Pt(10)
    for run in div.runs:
        _set_font(run, size=8, color=(209, 213, 219))

    # Date + label on same line
    meta_p = doc.add_paragraph()
    run_date = meta_p.add_run(datetime.now().strftime('%B %d, %Y'))
    _set_font(run_date, size=10, color=(107, 114, 128))
    run_label = meta_p.add_run(f'    {title_label}')
    _set_font(run_label, size=10, bold=True, color=(30, 58, 95))
    meta_p.paragraph_format.space_after = Pt(12)

    # Salutation
    sal_p = doc.add_paragraph(tailored.get('salutation', 'Dear Hiring Team,'))
    for run in sal_p.runs:
        _set_font(run, size=11)
    sal_p.paragraph_format.space_after = Pt(8)

    # Body paragraphs
    for para_key in para_keys:
        text = tailored.get(para_key, '')
        if text:
            p = doc.add_paragraph(text)
            for run in p.runs:
                _set_font(run, size=11)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = Pt(14)

    # Austria work authorization (cover letter only)
    if auth_paragraph:
        auth_text = (
            "I currently hold legal residence in Austria and am fully eligible "
            "to live and work locally without any sponsorship requirement."
        )
        p = doc.add_paragraph(auth_text)
        for run in p.runs:
            _set_font(run, size=11)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = Pt(14)

    # Sign off
    sign_p = doc.add_paragraph(tailored.get('sign_off', 'Kind regards,\nEmmanuel Rodríguez'))
    sign_p.paragraph_format.space_before = Pt(6)
    for run in sign_p.runs:
        _set_font(run, size=11, bold=True)

    fname_map = {
        'COVER LETTER': 'Emmanuel_Rodriguez_Cover_Letter.docx',
        'LETTER OF MOTIVATION': 'Emmanuel_Rodriguez_Motivation_Letter.docx',
    }
    fname = fname_map.get(title_label, f'Emmanuel_Rodriguez_{fname_prefix}.docx')
    path = os.path.join(OUTPUT_DIR, fname)
    doc.save(path)
    print(f"[generator] {title_label} saved: {path}")
    return path


def _build_cl_docx(tailored: dict, job: dict) -> str:
    return _build_letter_docx(
        tailored, job,
        para_keys=['opening', 'body', 'closing'],
        title_label='COVER LETTER',
        fname_prefix='CoverLetter',
        auth_paragraph=True,
    )


def _build_ml_docx(tailored: dict, job: dict) -> str:
    return _build_letter_docx(
        tailored, job,
        para_keys=['opening', 'motivation_field', 'closing'],
        title_label='LETTER OF MOTIVATION',
        fname_prefix='MotivationLetter',
        auth_paragraph=False,
    )


def generate_and_send(job: dict) -> tuple[str, str, str]:
    """
    Generate tailored CV + cover letter + motivation letter for a job.
    Saves as PDF (falls back to docx), emails all 3 to Emmanuel.
    Returns (cv_path, cl_path, ml_path).
    """
    print(f"[generator] Generating for: {job['title']} @ {job['company']}")

    try:
        tailored_cv = _tailor_cv(job)
    except Exception as e:
        print(f"[generator] Claude CV error: {e} — using master template")
        tailored_cv = {
            'summary': SUMMARY,
            'core_competencies': CORE_COMPETENCIES,
            'experience': EXPERIENCE,
        }

    try:
        tailored_cl = _tailor_cover_letter(job)
    except Exception as e:
        print(f"[generator] Claude CL error: {e} — using master template")
        tailored_cl = {
            'salutation': 'Dear Hiring Team,',
            'opening': STRUCTURE['opening'].replace('[ROLE]', job['title']).replace('[COMPANY]', job['company']),
            'body': '',
            'company_fit': '',
            'closing': STRUCTURE['closing'].replace('[COMPANY]', job['company']).replace('[KEY_AREA]', 'Consumer Insights & Commercial Strategy'),
            'sign_off': STRUCTURE['sign_off'],
        }

    try:
        tailored_ml = _tailor_motivation_letter(job)
    except Exception as e:
        print(f"[generator] Claude ML error: {e} — using fallback")
        tailored_ml = {
            'salutation': 'Dear Hiring Team,',
            'opening': f"The {job['title']} role at {job['company']} is a natural fit for where I am in my career and what I want to build next.",
            'motivation_field': 'Spending a decade in Consumer Insights and Commercial Strategy has reinforced one conviction: the sharpest competitive advantage in FMCG is the ability to translate data into decisions at speed. That is the work I do best and the work I want to keep doing at a higher level.',
            'closing': f"I am based in Austria with full work eligibility and ready to contribute from day one. I would welcome the chance to discuss how my background aligns with {job['company']}'s direction.",
            'sign_off': 'Kind regards,\nEmmanuel Rodríguez',
        }

    cv_docx = _build_cv_docx(tailored_cv, job)
    cl_docx = _build_cl_docx(tailored_cl, job)
    ml_docx = _build_ml_docx(tailored_ml, job)

    # Convert ALL three to PDF — must all succeed or all fall back to DOCX
    cv_pdf  = _docx_to_pdf(cv_docx)
    cl_pdf  = _docx_to_pdf(cl_docx)
    ml_pdf  = _docx_to_pdf(ml_docx)

    if cv_pdf and cl_pdf and ml_pdf:
        cv_path, cl_path, ml_path = cv_pdf, cl_pdf, ml_pdf
        print("[generator] All 3 documents converted to PDF successfully.")
    else:
        # At least one failed — send whatever was produced, but log clearly
        cv_path  = cv_pdf  or cv_docx
        cl_path  = cl_pdf  or cl_docx
        ml_path  = ml_pdf  or ml_docx
        missing = [n for n, p in [('CV', cv_pdf), ('CL', cl_pdf), ('ML', ml_pdf)] if not p]
        print(f"[generator] WARNING: PDF conversion failed for: {missing} — sending DOCX fallback for those.")

    # Save PDF copy + update Excel tracker
    try:
        record_sent_cv(job, cv_path)
    except Exception as e:
        print(f"[generator] Tracker error (non-fatal): {e}")

    # Send email with CV, cover letter, motivation letter and Excel tracker attached
    excel = EXCEL_PATH if os.path.exists(EXCEL_PATH) else None
    send_manual_package(job, cv_path, cl_path, ml_path=ml_path, excel_path=excel)
    update_status(job['id'], 'sent')

    return cv_path, cl_path, ml_path
